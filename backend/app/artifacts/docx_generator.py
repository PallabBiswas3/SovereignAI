from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


class DocxGenerator:
    def generate_report(
        self, output: Path, title: str, sections: list[dict[str, object]],
        sources: list[dict[str, object]] | None = None,
    ) -> Path:
        output.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        document.add_heading(title, level=0)
        for section in sections:
            document.add_heading(str(section.get("heading", "Section")), level=1)
            content = section.get("content", "")
            if isinstance(content, list):
                for item in content:
                    document.add_paragraph(str(item), style="List Bullet")
            else:
                document.add_paragraph(str(content))
        if sources:
            document.add_heading("Sources", level=1)
            for source in sources:
                document.add_paragraph(self._citation(source), style="List Bullet")
        document.save(output)
        return output

    def generate_approval_note(
        self,
        output: Path,
        title: str,
        subject: str,
        findings: list[dict[str, object]],
        recommendation: str,
        sources: list[dict[str, object]],
    ) -> Path:
        output.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        section = document.sections[0]
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        heading = document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph = document.add_paragraph()
        paragraph.add_run("Subject: ").bold = True
        paragraph.add_run(subject)
        document.add_heading("Purpose", level=1)
        document.add_paragraph("To record the evidence-based disposition of the inspected equipment against the applicable internal maintenance standard.")
        document.add_heading("Assessment", level=1)
        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        for cell, label in zip(table.rows[0].cells, ["Parameter", "Observed", "Allowed", "Status", "Evidence"]):
            cell.text = label
        for finding in findings:
            cells = table.add_row().cells
            cells[0].text = str(finding["parameter"])
            cells[1].text = str(finding["observed"])
            cells[2].text = str(finding["allowed"])
            cells[3].text = str(finding["status"])
            source = finding.get("source", {})
            cells[4].text = self._citation(source) if isinstance(source, dict) else ""
        document.add_heading("Recommendation", level=1)
        document.add_paragraph(recommendation)
        document.add_heading("Sources", level=1)
        for source in sources:
            document.add_paragraph(self._citation(source), style="List Bullet")
        warning = document.add_paragraph("Human authorization is required before operational action. This prototype supports decision-making; it does not replace engineering authority.")
        warning.runs[0].italic = True
        warning.runs[0].font.size = Pt(9)
        document.save(output)
        return output

    @staticmethod
    def _citation(source: dict[str, object]) -> str:
        parts = [str(source.get("file", "Unknown source"))]
        if source.get("page") is not None:
            parts.append(f"page {source['page']}")
        if source.get("section"):
            parts.append(f"section {source['section']}")
        return ", ".join(parts)
